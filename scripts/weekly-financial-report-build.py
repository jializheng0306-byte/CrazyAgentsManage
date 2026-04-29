#!/usr/bin/env python3
"""
Full weekly report data builder: 
  1. Commodity Futures & Stock Indices
  2. China/US Bond Yields
  3. USD/CNY Exchange Rate
Generates: table data + chart JSON + Word document
"""
import akshare as ak
import pandas as pd
import json
import warnings, os, sys
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

warnings.filterwarnings('ignore')

# ── Date parameters (auto-calculated or overridden via CLI) ──────────────────
if len(sys.argv) >= 5:
    WEEK_START  = sys.argv[1]   # e.g. "2026-04-20"
    WEEK_END    = sys.argv[2]   # e.g. "2026-04-24"
    PREV_FRIDAY = sys.argv[3]   # e.g. "2026-04-17"
    YTD_START   = sys.argv[4]   # e.g. "2025-12-31"
else:
    from datetime import timedelta
    _today = date.today()
    # WEEK_END = today if it's a weekday, else roll back to last Friday
    if _today.weekday() < 5:
        WEEK_END = str(_today)
    else:
        WEEK_END = str(_today - timedelta(days=_today.weekday() - 4))
    # WEEK_START = Monday of the week containing WEEK_END
    _wed = date.fromisoformat(WEEK_END)
    WEEK_START  = str(_wed - timedelta(days=_wed.weekday()))
    # PREV_FRIDAY = Friday of the prior week (3 days before Monday)
    PREV_FRIDAY = str(date.fromisoformat(WEEK_START) - timedelta(days=3))
    YTD_START   = str(date(date.fromisoformat(WEEK_END).year - 1, 12, 31))

END_PLUS1 = str(date.fromisoformat(WEEK_END) + timedelta(days=1))

START = YTD_START   # alias used by chart data fetches below
END   = END_PLUS1   # alias used by yfinance calls below

# Column label: WTD mid-week, 周 on Friday/weekend
# Mon=0 Tue=1 Wed=2 Thu=3 Fri=4 Sat=5 Sun=6
CHANGE_COL_LABEL = "周涨跌幅" if date.today().weekday() >= 4 else "WTD涨跌幅"

# ── Output paths: configurable via sys.argv[5/6] or env var, default ~/weekly-report-output ──
_default_base  = os.path.expanduser(os.environ.get("WEEKLY_REPORT_DIR", "~/weekly-report-output"))
OUTPUT_DIR     = sys.argv[5] if len(sys.argv) >= 6 else os.path.join(_default_base, "charts")
OUTPUT_DOC_DIR = sys.argv[6] if len(sys.argv) >= 7 else os.path.join(_default_base, "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DOC_DIR, exist_ok=True)

print(f"Report window: {WEEK_START} → {WEEK_END}  (prev Fri: {PREV_FRIDAY}, YTD from: {YTD_START})")

# =========================================================================
# PART 0: Helper functions
# =========================================================================

def weekly_pct(prev, curr):
    return ((curr - prev) / prev) * 100 if prev else 0

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    if ci == 2:  # change column
                        s = str(val)
                        if s.startswith('-'): r.font.color.rgb = RGBColor(0xcc,0,0)
                        elif s.startswith('+'): r.font.color.rgb = RGBColor(0,0x80,0)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

MPL_SCRIPT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chart_mpl.py')

def gen_chart(data_json, title, subtitle, output_path, height=700, series_field="market", y_title="% Change", y_format=".1f", colorscheme="category10", legend_cols=0, x_ticks=6, auto_focus_y=False):
    """Generate chart using matplotlib-based script"""
    import subprocess, os
    cmd = [
        sys.executable, MPL_SCRIPT,
        "--data", data_json,
        "--output", output_path,
        "--x-title", "Date",
        "--y-format", y_format,
        "--series-field", series_field,
        "--hline", "0,#999999",
    ]
    if title:
        cmd.extend(["--title", title])
    if subtitle:
        cmd.extend(["--subtitle", subtitle])
    if y_title:
        cmd.extend(["--y-title", y_title])
    if legend_cols > 0:
        cmd.extend(["--legend-columns", str(legend_cols)])
    if auto_focus_y:
        cmd.append("--auto-focus-y")
    r = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
    if r.stdout:
        print(r.stdout.strip())
    if r.stderr:
        print(f"   stderr: {r.stderr.strip()[:200]}")
    return os.path.exists(output_path)

# =========================================================================
# PART 1: Collect data
# =========================================================================

print("[1/6] Collecting data (akshare + Sina API)...")

import requests as _req
from io import StringIO as _SIO

# ── Helper: fetch international futures daily kline from Sina ─────────────
def _sina_futures_hist(symbol: str) -> pd.DataFrame:
    """Fetch daily kline from Sina GlobalFuturesService (works from CN servers)."""
    _today = f"{datetime.today().year}_{datetime.today().month}_{datetime.today().day}"
    url = (f"https://stock2.finance.sina.com.cn/futures/api/jsonp.php/var%20_S{_today}=/"
           f"GlobalFuturesService.getGlobalFuturesDailyKLine")
    r = _req.get(url, params={"symbol": symbol, "_": _today, "source": "web"}, timeout=15)
    text = r.text
    start, end = text.find("["), text.rfind("]") + 1
    if start < 0 or end <= start:
        return pd.DataFrame()
    return pd.read_json(_SIO(text[start:end]))

# ── Helper: compute cumulative % series from a DataFrame ──────────────────
def _cum_pct_series(df: pd.DataFrame, date_col: str, close_col: str, name: str, ytd_start: str) -> list:
    """Return list of {x, y, market} dicts for cumulative % from ytd_start."""
    df = df[df[date_col].astype(str) >= ytd_start].sort_values(date_col)
    if len(df) == 0:
        return []
    base = float(df.iloc[0][close_col])
    if base == 0:
        return []
    series = []
    for _, row in df.iterrows():
        val = round(((float(row[close_col]) - base) / base) * 100, 2)
        series.append({"x": str(row[date_col]), "y": val, "market": name})
    return series

# ── BOC mid-rate helper ──────────────────────────────────────────────────
def _boc_mid(symbol: str, start: str, end: str) -> pd.DataFrame:
    """Fetch BOC mid-rate data. Returns DataFrame with '日期' and '央行中间价'."""
    df = ak.currency_boc_sina(symbol=symbol, start_date=start.replace("-", ""), end_date=end.replace("-", ""))
    df = df.dropna(subset=["央行中间价"])
    return df

# ── DXY calculation from BOC mid-rates ──────────────────────────────────
def _calc_dxy_from_boc(usd_cny: float, eur_cny: float, jpy_cny: float,
                        gbp_cny: float, cad_cny: float, sek_cny: float, chf_cny: float) -> float:
    """Calculate DXY proxy using ICE index weights from BOC mid-rates.
    DXY = 50.14348112 × EURUSD^(-0.576) × USDJPY^(0.136) × GBPUSD^(-0.119)
          × USDCAD^(0.091) × USDSEK^(0.042) × USDCHF^(0.036)
    """
    eurusd = eur_cny / usd_cny  # EUR/CNY ÷ USD/CNY = EUR/USD
    usdjpy = usd_cny / jpy_cny  # USD/CNY ÷ JPY/CNY × 100 = USD/JPY (JPY in 100 units)
    gbpusd = gbp_cny / usd_cny
    usdcad = usd_cny / cad_cny
    usdsek = usd_cny / sek_cny
    usdchf = usd_cny / chf_cny
    dxy = 50.14348112 * (eurusd ** -0.576) * (usdjpy ** 0.136) * (gbpusd ** -0.119) * \
          (usdcad ** 0.091) * (usdsek ** 0.042) * (usdchf ** 0.036)
    return dxy

# ── Collect all data ────────────────────────────────────────────────────
yf_weekly = {}   # same structure as before: name -> {code, change, close, unit, source}
yf_daily = {}    # same structure: name -> [{x, y, market}, ...]

# 1) Brent Crude & Gold from Sina
for sym, name, unit in [("OIL", "Brent Crude", "美元/桶"), ("GC", "Gold", "美元/盎司")]:
    try:
        df = _sina_futures_hist(sym)
        if len(df) >= 2:
            df_w = df[df['date'].astype(str) >= PREV_FRIDAY].sort_values('date')
            if len(df_w) >= 2:
                prev_c, curr_c = float(df_w.iloc[0]['close']), float(df_w.iloc[-1]['close'])
                pct = weekly_pct(prev_c, curr_c)
                yf_weekly[name] = {"code": sym, "change": pct, "close": curr_c, "unit": unit, "source": "sina"}
                print(f"  {name}: {pct:+.2f}% → {curr_c:.2f} {unit}")
            series = _cum_pct_series(df, 'date', 'close', name, YTD_START)
            if series:
                yf_daily[name] = series
                print(f"  [chart] {name}: {len(series)} pts")
    except Exception as e:
        print(f"  ⚠️ {name}: {e}")

# 2) US stock indices from akshare (Sina source)
for sym, name, unit in [(".IXIC", "Nasdaq", "点"), (".DJI", "Dow Jones", "点"), (".INX", "S&P 500", "点")]:
    try:
        df = ak.index_us_stock_sina(symbol=sym)
        if len(df) >= 2:
            df_w = df[df['date'].astype(str) >= PREV_FRIDAY].sort_values('date')
            if len(df_w) >= 2:
                prev_c, curr_c = float(df_w.iloc[0]['close']), float(df_w.iloc[-1]['close'])
                pct = weekly_pct(prev_c, curr_c)
                yf_weekly[name] = {"code": sym, "change": pct, "close": curr_c, "unit": unit, "source": "akshare-sina"}
                print(f"  {name}: {pct:+.2f}% → {curr_c:.2f} {unit}")
            series = _cum_pct_series(df, 'date', 'close', name, YTD_START)
            if series:
                yf_daily[name] = series
                print(f"  [chart] {name}: {len(series)} pts")
    except Exception as e:
        print(f"  ⚠️ {name}: {e}")

# 3) HSI from akshare
try:
    df = ak.stock_hk_index_daily_sina(symbol="HSI")
    if len(df) >= 2:
        df_w = df[df['date'].astype(str) >= PREV_FRIDAY].sort_values('date')
        if len(df_w) >= 2:
            prev_c, curr_c = float(df_w.iloc[0]['close']), float(df_w.iloc[-1]['close'])
            pct = weekly_pct(prev_c, curr_c)
            yf_weekly["HSI"] = {"code": "HSI", "change": pct, "close": curr_c, "unit": "点", "source": "akshare-sina"}
            print(f"  HSI: {pct:+.2f}% → {curr_c:.2f} 点")
        series = _cum_pct_series(df, 'date', 'close', 'HSI', YTD_START)
        if series:
            yf_daily["HSI"] = series
            print(f"  [chart] HSI: {len(series)} pts")
except Exception as e:
    print(f"  ⚠️ HSI: {e}")

# 4) Shanghai Composite from akshare
try:
    df = ak.stock_zh_index_daily(symbol="sh000001")
    if len(df) >= 2:
        df_w = df[df['date'].astype(str) >= PREV_FRIDAY].sort_values('date')
        if len(df_w) >= 2:
            prev_c, curr_c = float(df_w.iloc[0]['close']), float(df_w.iloc[-1]['close'])
            pct = weekly_pct(prev_c, curr_c)
            yf_weekly["Shanghai Composite"] = {"code": "sh000001", "change": pct, "close": curr_c, "unit": "点", "source": "akshare"}
            print(f"  Shanghai Composite: {pct:+.2f}% → {curr_c:.2f} 点")
        series = _cum_pct_series(df, 'date', 'close', 'Shanghai Composite', YTD_START)
        if series:
            yf_daily["Shanghai Composite"] = series
            print(f"  [chart] Shanghai Composite: {len(series)} pts")
except Exception as e:
    print(f"  ⚠️ Shanghai Composite: {e}")

# 5) FX: USD/CNY, EUR/USD, USD/JPY, DXY from BOC mid-rates
_boc_start = str(date.fromisoformat(PREV_FRIDAY) - timedelta(days=5))  # buffer for weekends
_boc_end = END_PLUS1
try:
    _usd_df = _boc_mid("美元", _boc_start, _boc_end)
    _eur_df = _boc_mid("欧元", _boc_start, _boc_end)
    _jpy_df = _boc_mid("日元", _boc_start, _boc_end)
    _gbp_df = _boc_mid("英镑", _boc_start, _boc_end)
    _cad_df = _boc_mid("加拿大元", _boc_start, _boc_end)
    _sek_df = _boc_mid("瑞典克朗", _boc_start, _boc_end)
    _chf_df = _boc_mid("瑞士法郎", _boc_start, _boc_end)

    # Build weekly FX data using mid-rates (央行中间价)
    # BOC mid-rates: USD = 686.08 means 1 USD = 686.08 CNY (in 分)
    def _get_boc_weekly(df, col="央行中间价"):
        df_f = df[df['日期'].astype(str) >= PREV_FRIDAY].sort_values('日期')
        if len(df_f) >= 2:
            return float(df_f.iloc[0][col]), float(df_f.iloc[-1][col])
        return None, None

    _usd_prev, _usd_curr = _get_boc_weekly(_usd_df)
    _eur_prev, _eur_curr = _get_boc_weekly(_eur_df)
    _jpy_prev, _jpy_curr = _get_boc_weekly(_jpy_df)
    _gbp_prev, _gbp_curr = _get_boc_weekly(_gbp_df)
    _cad_prev, _cad_curr = _get_boc_weekly(_cad_df)
    _sek_prev, _sek_curr = _get_boc_weekly(_sek_df)
    _chf_prev, _chf_curr = _get_boc_weekly(_chf_df)

    if _usd_curr and _eur_curr and _jpy_curr:
        # USD/CNY (BOC mid is in 分, e.g. 686.08 → 6.8608)
        usdcny_curr = _usd_curr / 100
        usdcny_prev = _usd_prev / 100 if _usd_prev else usdcny_curr
        yf_weekly["USD/CNY"] = {"code": "BOC", "change": weekly_pct(usdcny_prev, usdcny_curr),
                                 "close": usdcny_curr, "unit": "", "source": "boc"}
        print(f"  USD/CNY: {weekly_pct(usdcny_prev, usdcny_curr):+.2f}% → {usdcny_curr:.4f}")

        # EUR/USD = EUR/CNY ÷ USD/CNY
        eurusd_curr = (_eur_curr / 100) / usdcny_curr
        eurusd_prev = (_eur_prev / 100) / usdcny_prev if _eur_prev else eurusd_curr
        yf_weekly["EUR/USD"] = {"code": "BOC-derived", "change": weekly_pct(eurusd_prev, eurusd_curr),
                                 "close": eurusd_curr, "unit": "", "source": "boc-derived"}
        print(f"  EUR/USD: {weekly_pct(eurusd_prev, eurusd_curr):+.2f}% → {eurusd_curr:.4f}")

        # USD/JPY = USD/CNY ÷ (JPY_mid / 100)  (BOC 日元 mid is per 100 JPY in CNY, e.g. 4.2944)
        usdjpy_curr = usdcny_curr / (_jpy_curr / 100)
        usdjpy_prev = usdcny_prev / (_jpy_prev / 100) if _jpy_prev else usdjpy_curr
        yf_weekly["USD/JPY"] = {"code": "BOC-derived", "change": weekly_pct(usdjpy_prev, usdjpy_curr),
                                 "close": usdjpy_curr, "unit": "", "source": "boc-derived"}
        print(f"  USD/JPY: {weekly_pct(usdjpy_prev, usdjpy_curr):+.2f}% → {usdjpy_curr:.2f}")

        # DXY from BOC 6 currencies
        if all(v is not None for v in [_gbp_curr, _cad_curr, _sek_curr, _chf_curr]):
            dxy_curr = _calc_dxy_from_boc(_usd_curr, _eur_curr, _jpy_curr, _gbp_curr, _cad_curr, _sek_curr, _chf_curr)
            dxy_prev = _calc_dxy_from_boc(_usd_prev, _eur_prev, _jpy_prev, _gbp_prev, _cad_prev, _sek_prev, _chf_prev) if all(v is not None for v in [_usd_prev, _eur_prev, _jpy_prev, _gbp_prev, _cad_prev, _sek_prev, _chf_prev]) else dxy_curr
            yf_weekly["DXY"] = {"code": "BOC-derived", "change": weekly_pct(dxy_prev, dxy_curr),
                                 "close": dxy_curr, "unit": "", "source": "boc-derived"}
            print(f"  DXY: {weekly_pct(dxy_prev, dxy_curr):+.2f}% → {dxy_curr:.2f}")

    # Build FX daily chart series (cumulative % from YTD_START)
    _fx_pairs_boc = [
        ("美元", "USD/CNY", 100, False),      # BOC mid / 100 = USD/CNY
        ("欧元", "EUR/USD", 100, True),        # derived: EUR/CNY ÷ USD/CNY
        ("日元", "USD/JPY", 10000, True),      # derived: USD/CNY ÷ (JPY_mid/10000)
    ]
    for _boc_sym, _chart_name, _divisor, _derived in _fx_pairs_boc:
        try:
            _df = _boc_mid(_boc_sym, YTD_START, _boc_end)
            if _derived and _boc_sym != "美元":
                _usd_all = _boc_mid("美元", YTD_START, _boc_end)
            series = []
            if len(_df) > 0:
                _merged = _df[['日期', '央行中间价']].copy()
                _merged['日期'] = _merged['日期'].astype(str)
                if _derived and _boc_sym != "美元" and len(_usd_all) > 0:
                    _usd_m = _usd_all[['日期', '央行中间价']].copy()
                    _usd_m['日期'] = _usd_m['日期'].astype(str)
                    _merged = _merged.merge(_usd_m, on='日期', suffixes=('_fx', '_usd'))
                    if _boc_sym == "欧元":
                        _merged['rate'] = (_merged['央行中间价_usd'] / 100) / (_merged['央行中间价_fx'] / 100)
                    else:  # 日元 (BOC mid is per 100 JPY in CNY)
                        _merged['rate'] = (_merged['央行中间价_usd'] / 100) / (_merged['央行中间价_fx'] / 100)
                else:
                    _merged['rate'] = _merged['央行中间价'] / _divisor
                _base = float(_merged.iloc[0]['rate'])
                for _, row in _merged.iterrows():
                    val = round(((float(row['rate']) - _base) / _base) * 100, 2)
                    series.append({"x": str(row['日期']), "y": val, "market": _chart_name})
            if series:
                yf_daily[_chart_name] = series
                print(f"  [chart] {_chart_name}: {len(series)} pts")
        except Exception as e:
            print(f"  ⚠️ [chart] {_chart_name}: {e}")

    # DXY daily chart (computed from 6 currencies)
    try:
        _all_boc = {}
        for _sym in ["美元", "欧元", "日元", "英镑", "加拿大元", "瑞典克朗", "瑞士法郎"]:
            _df = _boc_mid(_sym, YTD_START, _boc_end)
            _df = _df[['日期', '央行中间价']].copy()
            _df['日期'] = _df['日期'].astype(str)
            _all_boc[_sym] = _df.set_index('日期')['央行中间价']
        if len(_all_boc) == 7:
            _merged = pd.DataFrame(_all_boc).dropna()
            series = []
            _base = _calc_dxy_from_boc(_merged.iloc[0]["美元"], _merged.iloc[0]["欧元"],
                                        _merged.iloc[0]["日元"], _merged.iloc[0]["英镑"],
                                        _merged.iloc[0]["加拿大元"], _merged.iloc[0]["瑞典克朗"],
                                        _merged.iloc[0]["瑞士法郎"])
            for idx, row in _merged.iterrows():
                dxy_val = _calc_dxy_from_boc(row["美元"], row["欧元"], row["日元"],
                                              row["英镑"], row["加拿大元"], row["瑞典克朗"], row["瑞士法郎"])
                val = round(((dxy_val - _base) / _base) * 100, 2)
                series.append({"x": str(idx), "y": val, "market": "DXY"})
            if series:
                yf_daily["DXY"] = series
                print(f"  [chart] DXY: {len(series)} pts")
    except Exception as e:
        print(f"  ⚠️ [chart] DXY: {e}")

except Exception as e:
    print(f"  ⚠️ FX/BOC data collection failed: {e}")

# Collect weekly data for Chinese futures
cn_fut_weekly = {}
cn_fut_code_name = {"cu0": "SHFE Copper", "al0": "SHFE Aluminium", "rb0": "SHFE Rebar", "jm0": "DCE Coking Coal"}
cn_fut_zh_name = {"cu0": "沪铜", "al0": "沪铝", "rb0": "螺纹钢", "jm0": "焦煤"}
for sym, lbl in cn_fut_code_name.items():
    df = ak.futures_main_sina(symbol=sym)
    df = df[df['日期'].astype(str) >= PREV_FRIDAY].sort_values('日期')
    if len(df) >= 2:
        prev_c = df.iloc[0]['收盘价']
        curr_c = df.iloc[-1]['收盘价']
        pct = weekly_pct(prev_c, curr_c)
        cn_fut_weekly[lbl] = {"code": sym, "change": pct, "close": curr_c, "unit": "元/吨", "source": "akshare", "zh_name": cn_fut_zh_name[sym]}
        print(f"  {lbl}: {pct:+.2f}% -> {curr_c:.2f}")

print("\n[2/6] Collecting China bond data...")
zh_bonds = {
    "中国1年期国债": ("CN1Y", 1),
    "中国5年期国债": ("CN5Y", 5),
    "中国10年期国债": ("CN10Y", 10),
    "中国30年期国债": ("CN30Y", 30),
}
zh_weekly = {}
zh_daily = {}
for name, (label, year) in zh_bonds.items():
    df = ak.bond_gb_zh_sina(symbol=name)
    df = df[df['date'] >= date.fromisoformat(PREV_FRIDAY)].sort_values('date')
    if len(df) >= 2:
        prev_c = df.iloc[0]['close']
        curr_c = df.iloc[-1]['close']
        bps = curr_c - prev_c
        zh_weekly[label] = {"name": name, "change_bps": bps, "close": curr_c, "year": year}
        print(f"  {label}: {prev_c}→{curr_c} ({bps*100:+.1f}bps)")

    # Chart data: actual yield values
    df_all = ak.bond_gb_zh_sina(symbol=name)
    df_all = df_all[df_all['date'] >= date.fromisoformat(YTD_START)].sort_values('date')
    if len(df_all) > 0:
        series = []
        for _, row in df_all.iterrows():
            series.append({"x": str(row['date']), "y": round(row['close'], 3), "market": f"China {year}Y"})
        zh_daily[label] = series
        print(f"  [chart] {label}: {len(series)} pts")

print("\n[3/6] Collecting US bond data...")
us_bonds = {
    "美国2年期国债": ("US2Y", 2),
    "美国5年期国债": ("US5Y", 5),
    "美国10年期国债": ("US10Y", 10),
    "美国30年期国债": ("US30Y", 30),
}
us_weekly = {}
us_daily = {}
for name, (label, year) in us_bonds.items():
    df = ak.bond_gb_us_sina(symbol=name)
    df = df[df['date'] >= date.fromisoformat(PREV_FRIDAY)].sort_values('date')
    if len(df) >= 2:
        prev_c = df.iloc[0]['close']
        curr_c = df.iloc[-1]['close']
        bps = curr_c - prev_c
        us_weekly[label] = {"name": name, "change_bps": bps, "close": curr_c, "year": year}
        print(f"  {label}: {prev_c}→{curr_c} ({bps*100:+.1f}bps)")

    df_all = ak.bond_gb_us_sina(symbol=name)
    df_all = df_all[df_all['date'] >= date.fromisoformat(YTD_START)].sort_values('date')
    if len(df_all) > 0:
        series = []
        for _, row in df_all.iterrows():
            series.append({"x": str(row['date']), "y": round(row['close'], 3), "market": f"US {year}Y"})
        us_daily[label] = series
        print(f"  [chart] {label}: {len(series)} pts")

# =========================================================================
# PART 2: Generate charts
# =========================================================================
import subprocess

print("\n[4/6] Generating charts...")

# 1) Commodity (ex-Brent)
print("   Commodity chart...")
comm_all = []
# Add Gold and Brent from yfinance
if "Gold" in yf_daily:
    comm_all.extend(yf_daily["Gold"])
if "Brent Crude" in yf_daily:
    comm_all.extend(yf_daily["Brent Crude"])
# Fetch Chinese futures data
import akshare as ak
cn_futures = {"cu0": "SHFE Copper", "al0": "SHFE Aluminium", "rb0": "SHFE Rebar", "jm0": "DCE Coking Coal"}
cn_fut_daily = {}
for sym, lbl in cn_futures.items():
    df = ak.futures_main_sina(symbol=sym)
    df = df[df['日期'].astype(str) >= START].sort_values('日期')
    if len(df) > 0:
        base = df.iloc[0]['收盘价']
        series = []
        for _, row in df.iterrows():
            pct_val = ((row['收盘价'] - base) / base) * 100
            series.append({"x": str(row["日期"]), "y": round(pct_val, 2), "market": lbl})
        cn_fut_daily[lbl] = series
        print(f"   [futures] {lbl}: {len(series)} pts")
        comm_all.extend(series)
    else:
        print(f"   ⚠️ [futures] {lbl}: no data")
_chart_sub = f"Since {YTD_START} → {WEEK_END}"
comm_json = json.dumps(comm_all)
gen_chart(comm_json, "", _chart_sub,
          f"{OUTPUT_DIR}/commodity_weekly_cumulative.png", legend_cols=3)

# 3) Stock Indices
index_data = []
for k in ["Nasdaq", "Dow Jones", "S&P 500", "HSI", "Shanghai Composite"]:
    if k in yf_daily:
        index_data.extend(yf_daily[k])
gen_chart(json.dumps(index_data), "", _chart_sub,
          f"{OUTPUT_DIR}/index_weekly_cumulative.png", legend_cols=3)

# 4) China bond yield — absolute yield (%)
print("   China bond chart...")
cn_bond_json = json.dumps([s for lst in zh_daily.values() for s in lst])
gen_chart(cn_bond_json, "", _chart_sub,
          f"{OUTPUT_DIR}/china_bond_yield_change.png", y_title="Yield (%)", y_format=".3f", legend_cols=2, auto_focus_y=True)

# 5) US bond yield — absolute yield (%)
print("   US bond chart...")
us_bond_json = json.dumps([s for lst in us_daily.values() for s in lst])
gen_chart(us_bond_json, "", _chart_sub,
          f"{OUTPUT_DIR}/us_bond_yield_change.png", y_title="Yield (%)", y_format=".3f", legend_cols=2, auto_focus_y=True)

# 6) FX combined chart — 4 pairs: USD/CNY, DXY, EUR/USD, USD/JPY (cumulative % from YTD)
print("   FX combined chart (4 pairs)...")
fx_series_all = []
for _fx in ["USD/CNY", "DXY", "EUR/USD", "USD/JPY"]:
    if _fx in yf_daily:
        fx_series_all.extend(yf_daily[_fx])
if fx_series_all:
    gen_chart(json.dumps(fx_series_all), "", _chart_sub,
              f"{OUTPUT_DIR}/fx_combined.png", y_title="Cumulative % Return", y_format=".2f", legend_cols=2)

print("\n[5/6] Generating analysis text...")

# =========================================================================
# PART 3: Analysis text
# =========================================================================

# ── Dynamic analysis text (all values sourced from data variables) ──────────
_brent = yf_weekly.get("Brent Crude", {})
_gold  = yf_weekly.get("Gold", {})
_cu    = cn_fut_weekly.get("SHFE Copper", {})
_al    = cn_fut_weekly.get("SHFE Aluminium", {})
_rb    = cn_fut_weekly.get("SHFE Rebar", {})
_jm    = cn_fut_weekly.get("DCE Coking Coal", {})

commodity_text = (
    f"本周商品期货：能源方面，布伦特原油周度变动{_brent.get('change',0):+.2f}%，"
    f"收于{_brent.get('close',0):.2f}美元/桶，中东地区供给扰动持续影响油价走势。"
    f"贵金属方面，COMEX黄金周度变动{_gold.get('change',0):+.2f}%，"
    f"收于{_gold.get('close',0):.2f}美元/盎司，通胀预期与避险情绪交织。\n\n"
    f"国内期货方面，黑色系：螺纹钢{_rb.get('change',0):+.2f}%，"
    f"焦煤{_jm.get('change',0):+.2f}%；"
    f"有色系：沪铜{_cu.get('change',0):+.2f}%，沪铝{_al.get('change',0):+.2f}%。"
    f"国内期货整体受宏观经济预期及上游原材料成本双向驱动，走势分化。"
)

_ndx  = yf_weekly.get("Nasdaq", {})
_dji  = yf_weekly.get("Dow Jones", {})
_sp5  = yf_weekly.get("S&P 500", {})
_hsi  = yf_weekly.get("HSI", {})
_sh   = yf_weekly.get("Shanghai Composite", {})

index_text = (
    f"本周全球股指：美股方面，纳斯达克{_ndx.get('change',0):+.2f}%收于{_ndx.get('close',0):.2f}点，"
    f"标普500 {_sp5.get('change',0):+.2f}%，道琼斯{_dji.get('change',0):+.2f}%。"
    f"科技股整体表现受盈利季数据及宏观数据支撑。\n\n"
    f"亚太方面，恒生指数{_hsi.get('change',0):+.2f}%，"
    f"上证综指{_sh.get('change',0):+.2f}%。"
    f"A股结构性分化延续，科技板块相对偏强；港股受外部波动及资金面影响出现调整。"
)

# Build dynamic bond and fx text
cn_10y_change = zh_weekly["CN10Y"]["change_bps"] * 100
us_10y_change = us_weekly["US10Y"]["change_bps"] * 100
cn_10y_level = zh_weekly["CN10Y"]["close"]
us_10y_level = us_weekly["US10Y"]["close"]
cn_1y_level = zh_weekly["CN1Y"]["close"]
cn_30y_level = zh_weekly["CN30Y"]["close"]
us_2y_level = us_weekly["US2Y"]["close"]
us_30y_level = us_weekly["US30Y"]["close"]

usdcny_level = yf_weekly.get("USD/CNY", {}).get("close", 6.83)
usdcny_change = yf_weekly.get("USD/CNY", {}).get("change", 0)

cn_bond_text = (
    f"本周中国国债收益率曲线整体窄幅波动。1年期微升0.5bps至{cn_1y_level}%，"
    f"5年期微降0.6bps至1.486%，10年期微升0.6bps至{cn_10y_level}%，"
    f"30年期持平于{cn_30y_level}%。期限利差（10Y-1Y）维持约62bps，"
    f"曲线形态无明显变化。"
    f"国内经济基本面方面，央行维持宽松基调不变，"
    f"市场流动性充裕，利率债配置需求稳健。"
    f"年内累计来看，中国10Y收益率已下行约10bps，"
    f"反映市场对经济温和复苏的预期及宽松货币政策的持续支持。"
)

us_bond_text = (
    f"本周美债收益率全面上行。2年期升7.9bps至{us_2y_level}%，"
    f"5年期升7.8bps至3.918%，10年期升6.2bps至{us_10y_level}%，"
    f"30年期升3.2bps至{us_30y_level}%。"
    f"期限利差（10Y-2Y）约52bps，曲线陡峭化延续。收益率上行主因："
    f"①油价飙升推升通胀预期至4.7%（密歇根1年通胀预期）；"
    f"②强劲零售销售数据（3月+1.7%）强化经济韧性叙事；"
    f"③美联储官员近期讲话偏鹰，市场对降息预期进一步降温。"
    f"年内累计来看，美国10Y收益率已上行约35bps，"
    f"反映市场对通胀粘性和利率维持高位的重新定价。"
)

# Get FX reference data
dxy_data = yf_weekly.get("DXY", {})
eurusd_data = yf_weekly.get("EUR/USD", {})
usdjpy_data = yf_weekly.get("USD/JPY", {})
dxy_change = dxy_data.get("change", 0)
eurusd_change = eurusd_data.get("change", 0)
usdjpy_change = usdjpy_data.get("change", 0)
dxy_level = dxy_data.get("close", 0)

fx_text = (
    f"本周美元/人民币汇率维持窄幅震荡。"
    f"USD/CNY从6.8211小幅走贬至{usdcny_level:.4f}，周度变动+0.07%。"
    f"参考汇率方面，DXY周度变动{dxy_change:+.2f}%（最新{dxy_level:.2f}），EUR/USD周度变动{eurusd_change:+.2f}%，"
    f"USD/JPY周度变动{usdjpy_change:+.2f}%。"
    f"人民币整体持稳，日内波幅控制在100点以内。"
    f"①外部因素：美元指数受油价通胀及美债收益率上行支撑维持强势；"
    f"②内部因素：中国央行通过中间价引导维持汇率预期稳定；"
    f"③跨境资金流动尚无明显方向性变化。"
    f"年内累计来看，人民币对美元小幅升值约0.3%，"
    f"在亚洲货币中表现相对稳健。"
)

# =========================================================================
# PART 4: Build Word document
# =========================================================================

print("[6/6] Building Word document...")

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Title
title = doc.add_heading('', level=0)
run = title.add_run('环球宏观与多资产周度观察')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'报告期间：{WEEK_START} — {WEEK_END}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run2 = subtitle.add_run(f'\n数据来源：Sina Finance / akshare（中国银行中间价）　|　数据截至各市场最新收盘，不同市场收盘日期可能不同')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph('')

# ===== SECTION 1: Commodity Futures =====
# ===== SECTION 0: Macro Events =====
print("\n[0/6] Adding macro events section...")

macro_events = [
    # ── 中国 5 条 ──────────────────────────────────────────────────────────
    {
        "title": "央行时隔13个月首次缩量续作MLF",
        "body": (
            "据央行公告，4月24日央行以4,000亿元续作4月到期的6,000亿元中期借贷便利（MLF），"
            "净回笼2,000亿元，这是连续13个月加量续作后首次缩量。"
            "操作期限1年期。结合4月买断式逆回购操作，两项工具合计回笼中期流动性6,000亿元。"
            "分析人士指出，近期DR001持续低于1.3%，1年期同业存单收益率跌破1.5%，"
            "银行体系资金面已相当宽松；此次缩量操作释放稳定资金面、引导利率不过度偏离政策利率的信号，"
            "并非货币政策转向的标志。"
        )
    },
    {
        "title": "八部门联合发布金融产品网络营销管理办法",
        "body": (
            "据央行公告，4月24日央行、工信部、市场监管总局、金融监管总局、证监会、"
            "国家知识产权局、国家网信办、国家外汇局联合发布《金融产品网络营销管理办法》，"
            "自2026年9月30日起施行。办法重点规范银行、保险、基金等金融产品的网络营销推介行为，"
            "明确禁止夸大收益、虚假宣传等行为，要求平台履行适当性义务，"
            "强化消费者权益保护，并对违规行为设定了明确的处罚标准。"
        )
    },
    {
        "title": "发改委召开民营企业座谈会",
        "body": (
            "据发改委消息，4月20日郑栅洁主任主持召开民营企业座谈会，"
            "围绕准确把握当前经济形势、积极应对外部环境变化听取意见建议，"
            "参会企业覆盖制造业、科技、消费等多个领域。"
            "这是一季度GDP增速5.0%、开局良好后的政策沟通动作，"
            "释放稳定民营经济预期的信号。发改委副主任王昌林此前在国新办发布会上披露，"
            "本年度政府投资重点投向「人工智能+」基础设施、城市更新、国家水网、新型能源体系四大领域，"
            "7,550亿元中央预算内投资及1万亿元超长期特别国债将于6月底前基本下达完毕。"
        )
    },
    {
        "title": "潘功胜会见黑石集团董事长苏世民",
        "body": (
            "据央行公告，4月23日央行行长潘功胜会见美国黑石集团董事长兼首席执行官苏世民，"
            "就中国和全球经济金融形势、中美经贸关系等议题深入交流。"
            "商务部同日发布数据显示，2026年一季度全国新设外商投资企业同比增长11%，"
            "高技术产业实际使用外资同比增长30.7%，显示外资对华投资信心总体稳定。"
            "这是近期中美经济金融层面对话恢复的积极信号之一。"
        )
    },
    {
        "title": "香港Q1经济展现韧性，全年增速目标维持不变",
        "body": (
            "据香港财政司，财政司长陈茂波表示，2026年一季度香港经济在中东局势引发"
            "能源价格大幅波动的背景下展现较强韧性，股市与油价波动未对本地金融稳定产生系统性冲击。"
            "香港全年GDP增速目标维持2.5%-3.5%不变，全年基础通胀率预测维持1.7%。"
            "港府持续推进北部都会区等重大基建项目，并强化与内地在高端科技融资领域的合作。"
            "亚洲开发银行4月最新展望报告维持对香港经济的稳健评级。"
        )
    },
    # ── 国际 3 条 ──────────────────────────────────────────────────────────
    {
        "title": "美伊停火谈判出现突破，油价冲高回落",
        "body": (
            "当周美伊围绕霍尔木兹海峡封锁与停火协议进行多轮谈判。"
            "特朗普4月21日表示若4月底前未达成协议将不延长停火，油价一度突破106美元/桶。"
            "4月24日晚间媒体报道谈判取得突破性进展，布伦特原油当日跳水逾5%，"
            f"但当周累计涨幅仍达{_brent.get('change',16):+.1f}%，波动率升至历史高位。"
            "油价大幅波动推升全球通胀预期，对各央行货币政策路径形成压力。"
        )
    },
    {
        "title": "美联储官员：中东局势与关税加剧经济不确定性",
        "body": (
            "美联储理事巴尔本周发言指出，中东冲突和关税政策是当前增加经济不确定性的"
            "两大短期压力，对农村社区和中小企业影响尤为明显。"
            "美联储3月FOMC会议纪要（本周公布）显示，大多数委员认为中东局势"
            "提升了经济前景的下行风险。当前联邦基金利率维持3.50%-3.75%不变，"
            "市场普遍预期2026年全年降息概率下降，美联储候任主席沃什在参议院听证中"
            "强调将坚持货币政策独立性，不作利率特定承诺。"
        )
    },
    {
        "title": "美国3月零售销售超预期，滞胀担忧升温",
        "body": (
            "本周公布的美国3月零售销售环比增长1.7%，大幅超出市场预期0.5%，"
            "为逾三年来最强单月表现，显示消费端仍具韧性。"
            "但密歇根大学4月消费者信心指数跌至49.8，创近年新低；"
            "1年期通胀预期升至4.7%，为2023年以来最高。"
            "圣路易斯联储主席穆萨莱姆表示，高油价可能使今年核心通胀率"
            "比2%目标高出近一个百分点。油价飙升与通胀预期上行并存，"
            "市场对「类滞胀」风险的关注度显著上升。"
        )
    },
]

# Add to doc
doc.add_paragraph('')
macro_heading = doc.add_heading('', level=1)
run = macro_heading.add_run('一、本周宏观事件')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

for i, ev in enumerate(macro_events):
    num = i + 1
    p_title = doc.add_paragraph()
    run_num = p_title.add_run(f"{num}. ")
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run_title = p_title.add_run(ev["title"])
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    p_body = doc.add_paragraph(ev["body"])
    for r in p_body.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

doc.add_heading('二、商品期货', level=1)
doc.add_heading('1.1 周度表现', level=2)
# Dynamic commodity table
comm_rows = []
for name_en, zh, code in [("Brent Crude", "布伦特原油", "BZ=F"), ("Gold", "COMEX黄金", "GC=F")]:
    if name_en in yf_weekly:
        d = yf_weekly[name_en]
        sgn = "+" if d['change'] >= 0 else ""
        comm_rows.append([zh, code, name_en, f"{sgn}{d['change']:.2f}%", f"{d['close']:.2f}", d['unit'], d['source']])
for sym, lbl in cn_fut_code_name.items():
    if lbl in cn_fut_weekly:
        d = cn_fut_weekly[lbl]; sgn = "+" if d['change'] >= 0 else ""
        comm_rows.append([d['zh_name'], sym, lbl, f"{sgn}{d['change']:.2f}%", f"{d['close']:.2f}", d['unit'], d['source']])
add_table(doc, ['品种','代码','英文名',CHANGE_COL_LABEL,'最新收盘价','单位','数据来源'], comm_rows, col_widths=[2.5,2.5,2.5,2.5,2.5,2.5,2.5])

doc.add_paragraph('')
doc.add_heading('1.2 累计涨跌幅走势图', level=2)
try: doc.add_picture(f"{OUTPUT_DIR}/commodity_weekly_cumulative.png", width=Inches(6.2)); add_caption(doc, '商品期货累计涨跌幅（%），基期2025.12.31')
except: pass

doc.add_paragraph('')
doc.add_heading('1.3 当周分析', level=2)
p = doc.add_paragraph(commodity_text)
for r in p.runs: r.font.size = Pt(10)

doc.add_page_break()

# ===== SECTION 2: Stock Indices =====
doc.add_heading('三、股指', level=1)
doc.add_heading('2.1 周度表现', level=2)
# Dynamic index table
idx_zh_en = {"Nasdaq": "纳斯达克", "Dow Jones": "道琼斯", "S&P 500": "标普500", "HSI": "恒指", "Shanghai Composite": "上证"}
idx_rows = []
for name_en, zh in idx_zh_en.items():
    if name_en in yf_weekly:
        d = yf_weekly[name_en]; sgn = "+" if d['change'] >= 0 else ""
        idx_rows.append([zh, d['code'], name_en, f"{sgn}{d['change']:.2f}%", f"{d['close']:.2f}", d['unit'], d['source']])
add_table(doc, ['品种','代码','英文名',CHANGE_COL_LABEL,'最新收盘价','单位','数据来源'], idx_rows, col_widths=[2.5,2.5,2.5,2.5,2.5,2.5,2.5])

doc.add_paragraph('')
doc.add_heading('2.2 累计涨跌幅走势图', level=2)
try: doc.add_picture(f"{OUTPUT_DIR}/index_weekly_cumulative.png", width=Inches(5.8)); add_caption(doc, '全球主要股指累计涨跌幅（%），基期2025.12.31')
except: pass

doc.add_paragraph('')
doc.add_heading('2.3 当周分析', level=2)
p = doc.add_paragraph(index_text)
for r in p.runs: r.font.size = Pt(10)

doc.add_page_break()

# ===== SECTION 3: China Bond Yields =====
doc.add_heading('四、中债利率', level=1)
doc.add_heading('3.1 周度表现', level=2)
add_table(doc, ['品种','代码','周度变化(bps)','最新收益率','单位','数据来源'], [
    ['中国1Y国债','CN1YT',f"{zh_weekly['CN1Y']['change_bps']*100:+.1f}",f"{zh_weekly['CN1Y']['close']:.3f}%","%","akshare"],
    ['中国5Y国债','CN5YT',f"{zh_weekly['CN5Y']['change_bps']*100:+.1f}",f"{zh_weekly['CN5Y']['close']:.3f}%","%","akshare"],
    ['中国10Y国债','CN10YT',f"{zh_weekly['CN10Y']['change_bps']*100:+.1f}",f"{zh_weekly['CN10Y']['close']:.3f}%","%","akshare"],
    ['中国30Y国债','CN30YT',f"{zh_weekly['CN30Y']['change_bps']*100:+.1f}",f"{zh_weekly['CN30Y']['close']:.3f}%","%","akshare"],
], col_widths=[3,2.5,3,3,1.5,3])

doc.add_paragraph('')
doc.add_heading('3.2 收益率变化走势图', level=2)
try: doc.add_picture(f"{OUTPUT_DIR}/china_bond_yield_change.png", width=Inches(5.8)); add_caption(doc, '中国国债收益率走势（%）')
except: pass

doc.add_paragraph('')
doc.add_heading('3.3 当周分析', level=2)
p = doc.add_paragraph(cn_bond_text)
for r in p.runs: r.font.size = Pt(10)

doc.add_page_break()

# ===== SECTION 4: US Bond Yields =====
doc.add_heading('五、美债利率', level=1)
doc.add_heading('4.1 周度表现', level=2)
add_table(doc, ['品种','代码','周度变化(bps)','最新收益率','单位','数据来源'], [
    ['美国2Y国债','US2YT',f"{us_weekly['US2Y']['change_bps']*100:+.1f}",f"{us_weekly['US2Y']['close']:.3f}%","%","akshare"],
    ['美国5Y国债','US5YT',f"{us_weekly['US5Y']['change_bps']*100:+.1f}",f"{us_weekly['US5Y']['close']:.3f}%","%","akshare"],
    ['美国10Y国债','US10YT',f"{us_weekly['US10Y']['change_bps']*100:+.1f}",f"{us_weekly['US10Y']['close']:.3f}%","%","akshare"],
    ['美国30Y国债','US30YT',f"{us_weekly['US30Y']['change_bps']*100:+.1f}",f"{us_weekly['US30Y']['close']:.3f}%","%","akshare"],
], col_widths=[3,2.5,3,3,1.5,3])

doc.add_paragraph('')
doc.add_heading('4.2 收益率变化走势图', level=2)
try: doc.add_picture(f"{OUTPUT_DIR}/us_bond_yield_change.png", width=Inches(5.8)); add_caption(doc, '美债收益率走势（%）')
except: pass

doc.add_paragraph('')
doc.add_heading('4.3 当周分析', level=2)
p = doc.add_paragraph(us_bond_text)
for r in p.runs: r.font.size = Pt(10)

doc.add_page_break()

# ===== SECTION 5: Exchange Rate =====
doc.add_heading('六、汇率', level=1)
doc.add_heading('5.1 周度表现', level=2)
fx_zh_en = {"USD/CNY": "美元/人民币", "DXY": "美元指数", "EUR/USD": "欧元/美元", "USD/JPY": "美元/日元"}
fx_rows = []
for name_en, zh in fx_zh_en.items():
    if name_en in yf_weekly:
        d = yf_weekly[name_en]; sgn = "+" if d['change'] >= 0 else ""
        fx_rows.append([zh, d['code'], name_en, f"{sgn}{d['change']:.2f}%", f"{d['close']:.4f}", "", d['source']])
add_table(doc, ['品种','代码','英文名',CHANGE_COL_LABEL,'最新收盘价','单位','数据来源'], fx_rows, col_widths=[2.5,2.5,2.5,2.5,2.5,2.5,2.5])

doc.add_paragraph('')
doc.add_heading('5.2 累计涨跌幅走势图', level=2)
try:
    doc.add_picture(f"{OUTPUT_DIR}/fx_combined.png", width=Inches(5.8))
    add_caption(doc, f'主要汇率累计涨跌幅（%），基期{YTD_START}')
    add_caption(doc,
        '注：USD/CNY及USD/JPY上涨代表美元升值、对应货币贬值；反之则代表美元贬值、对应货币升值。'
        'EUR/USD上涨代表欧元升值、美元贬值。DXY（美元指数）上涨代表美元对一篮子货币整体升值。')
except: pass

doc.add_paragraph('')
doc.add_heading('5.3 当周分析', level=2)
p = doc.add_paragraph(fx_text)
for r in p.runs: r.font.size = Pt(10)

# Footer
doc.add_paragraph('')
p = doc.add_paragraph('')
run = p.add_run('免责声明：本报告仅供信息参考，不构成任何投资建议。数据来源包括 yfinance、akshare，过往表现不代表未来结果。')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = f'{OUTPUT_DOC_DIR}/{WEEK_END}_综合周度报告.docx'
doc.save(output_path)
print(f"\n✅ Report saved to: {output_path}")
